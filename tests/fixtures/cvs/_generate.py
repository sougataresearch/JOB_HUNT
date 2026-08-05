"""One-off script to (re)generate the fixture CVs under tests/fixtures/cvs/.

Not a test itself (prefixed with ``_`` so pytest doesn't collect it,
and not imported anywhere in the test suite). Run manually with:

    python tests/fixtures/cvs/_generate.py

Requires ``reportlab`` (dev dependency, fixture-generation only) and
``python-docx`` (already a runtime dependency for DOCXParser).

Three personas, each with the same content across all three formats
(tasks.md T5.1: "3+ fixture CVs per format"), so the Resume Analysis
Agent's golden-file tests (T5.2) can expect the same extracted profile
regardless of source format:

- jane_doe_complete: a fully detailed CV.
- sparse_partial: deliberately minimal, to exercise "explicit not
  found, never guessed" (phases.md Phase 5 acceptance criteria).
- alex_kim_standard: a second, differently-worded complete CV.
"""

from __future__ import annotations

from pathlib import Path
from typing import TypedDict

import docx
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).resolve().parent


class _Persona(TypedDict):
    name: str
    lines: list[str]


_JANE_DOE: _Persona = {
    "name": "jane_doe_complete",
    "lines": [
        "# Jane Doe",
        "jane.doe@example.com | +1-555-0101 | Seattle, WA",
        "",
        "## Summary",
        "Backend engineer with 6 years of experience building distributed "
        "systems in Python and Go.",
        "",
        "## Experience",
        "Senior Software Engineer, Acme Corp (2021-2024)",
        "- Led migration of the payments service from a monolith to "
        "independently deployed microservices.",
        "- Reduced p99 API latency by 40 percent by introducing request-level caching.",
        "",
        "Software Engineer, Beta Systems (2018-2021)",
        "- Built the initial version of the internal job-scheduling platform.",
        "",
        "## Education",
        "B.S. Computer Science, State University (2014-2018)",
        "",
        "## Skills",
        "Python, Go, PostgreSQL, Kubernetes, AWS",
        "",
        "## Certifications",
        "AWS Certified Solutions Architect",
    ],
}

_SPARSE: _Persona = {
    "name": "sparse_partial",
    "lines": [
        "# A. Candidate",
        "",
        "## Experience",
        "Worked at a small startup for about two years doing various engineering tasks.",
    ],
}

_ALEX_KIM: _Persona = {
    "name": "alex_kim_standard",
    "lines": [
        "# Alex Kim",
        "alex.kim@example.com | Austin, TX",
        "",
        "## Summary",
        "Data engineer focused on ETL pipelines and analytics infrastructure.",
        "",
        "## Experience",
        "Data Engineer, Gamma Analytics (2019-2024)",
        "- Designed and maintained Airflow pipelines processing 2TB of data daily.",
        "- Mentored two junior engineers.",
        "",
        "## Education",
        "M.S. Data Science, Tech Institute (2017-2019)",
        "B.A. Mathematics, Liberal Arts College (2013-2017)",
        "",
        "## Skills",
        "Python, SQL, Airflow, Spark, dbt",
    ],
}

_PERSONAS: list[_Persona] = [_JANE_DOE, _SPARSE, _ALEX_KIM]


def _write_markdown(persona: _Persona) -> None:
    path = FIXTURES_DIR / f"{persona['name']}.md"
    path.write_text("\n".join(persona["lines"]) + "\n", encoding="utf-8")


def _write_docx(persona: _Persona) -> None:
    document = docx.Document()
    for line in persona["lines"]:
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line:
            document.add_paragraph(line)
    document.save(str(FIXTURES_DIR / f"{persona['name']}.docx"))


def _write_pdf(persona: _Persona) -> None:
    path = FIXTURES_DIR / f"{persona['name']}.pdf"
    pdf_canvas = canvas.Canvas(str(path), pagesize=LETTER)
    _, height = LETTER
    y = height - 50
    for line in persona["lines"]:
        if line.startswith("## "):
            text = line[3:]
        elif line.startswith("# "):
            text = line[2:]
        else:
            text = line
        pdf_canvas.drawString(50, y, text)
        y -= 16
        if y < 50:
            pdf_canvas.showPage()
            y = height - 50
    pdf_canvas.save()


def main() -> None:
    """Generate all three formats for every persona."""
    for persona in _PERSONAS:
        _write_markdown(persona)
        _write_docx(persona)
        _write_pdf(persona)
        print(f"Generated fixtures for {persona['name']}")


if __name__ == "__main__":
    main()
