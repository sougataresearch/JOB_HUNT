"""DOCX CV parser (api.md §1 CV Parser API).

Uses ``python-docx`` for paragraph extraction.
"""

from __future__ import annotations

from pathlib import Path

import docx

from jobhunt_core.documents.parsers.sectioning import split_plaintext_sections
from jobhunt_core.schemas.document import ParsedDocument


class DOCXParser:
    """Parses ``.docx`` CV files."""

    def supports(self, file_path: Path) -> bool:
        """True for the ``.docx`` extension."""
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path) -> ParsedDocument:
        """Extract paragraph text and split into best-effort sections."""
        document = docx.Document(str(file_path))
        raw_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ParsedDocument(
            raw_text=raw_text,
            sections=split_plaintext_sections(raw_text),
            source_format="docx",
        )
